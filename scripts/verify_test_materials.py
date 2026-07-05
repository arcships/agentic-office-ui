from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from PIL import Image
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
SAMPLES = ROOT / "apps" / "demo" / "public" / "samples"

valid_docx = [
    "demo.docx",
    "legal-contract.docx",
    "invoice-table.docx",
    "report-with-image.docx",
    "chinese-mixed.docx",
]
valid_pdf = [
    "sample.pdf",
    "scanned-invoice.pdf",
    "rotated-pages.pdf",
    "large-contract.pdf",
]
valid_xlsx = [
    "financial-model.xlsx",
    "sales-table.xlsx",
    "charts-images.xlsx",
    "large-grid.xlsx",
]
valid_png = ["invoice.png", "contract-page.png"]
valid_json = ["field-citations.json", "ocr-layout.json", "manifest.json"]
corrupted = ["corrupted.docx", "corrupted.pdf", "corrupted.xlsx"]

report = {"docx": {}, "pdf": {}, "xlsx": {}, "png": {}, "json": {}, "corrupted": {}}

for name in valid_docx:
    path = SAMPLES / name
    doc = Document(path)
    report["docx"][name] = {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "bytes": path.stat().st_size,
    }

for name in valid_pdf:
    path = SAMPLES / name
    reader = PdfReader(str(path))
    report["pdf"][name] = {
        "pages": len(reader.pages),
        "bytes": path.stat().st_size,
    }

for name in valid_xlsx:
    path = SAMPLES / name
    wb = load_workbook(path, data_only=False, read_only=False)
    report["xlsx"][name] = {
        "sheets": wb.sheetnames,
        "bytes": path.stat().st_size,
    }
    wb.close()

for name in valid_png:
    path = SAMPLES / name
    with Image.open(path) as img:
        report["png"][name] = {
            "size": img.size,
            "mode": img.mode,
            "bytes": path.stat().st_size,
        }

for name in valid_json:
    path = SAMPLES / name
    data = json.loads(path.read_text(encoding="utf-8"))
    report["json"][name] = {
        "type": type(data).__name__,
        "bytes": path.stat().st_size,
    }

for name in corrupted:
    path = SAMPLES / name
    is_zip = zipfile.is_zipfile(path)
    report["corrupted"][name] = {
        "is_zip": is_zip,
        "bytes": path.stat().st_size,
    }
    if is_zip:
        raise AssertionError(f"{name} should be a negative corrupted fixture, but is a zip")

print(json.dumps(report, ensure_ascii=False, indent=2))
