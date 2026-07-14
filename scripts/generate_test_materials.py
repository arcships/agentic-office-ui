from __future__ import annotations

import json
import random
import re
import zipfile
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.drawing.image import Image as XlsxImage
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.comments import Comment
from openpyxl.formatting.rule import ColorScaleRule, DataBarRule, IconSetRule
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table as XlsxTable, TableStyleInfo
from PIL import Image, ImageDraw, ImageFont
from reportlab import rl_config
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
SAMPLES = ROOT / "apps" / "demo" / "public" / "samples"
SAMPLES.mkdir(parents=True, exist_ok=True)
random.seed(42)
rl_config.invariant = 1

FIXED_OOXML_TIMESTAMP = "2020-01-01T00:00:00Z"
FIXED_ZIP_TIMESTAMP = (2020, 1, 1, 0, 0, 0)


def normalize_ooxml_zip(path: Path):
    """Make OOXML bytes reproducible across repeated local and CI runs."""
    temporary = path.with_suffix(f"{path.suffix}.deterministic")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for original in sorted(source.infolist(), key=lambda item: item.filename):
            data = source.read(original.filename)
            if original.filename == "docProps/core.xml":
                text = data.decode("utf-8")
                text = re.sub(
                    r"(<dcterms:(?:created|modified)[^>]*>).*?(</dcterms:(?:created|modified)>)",
                    rf"\g<1>{FIXED_OOXML_TIMESTAMP}\g<2>",
                    text,
                )
                data = text.encode("utf-8")
            entry = zipfile.ZipInfo(original.filename, FIXED_ZIP_TIMESTAMP)
            entry.compress_type = zipfile.ZIP_DEFLATED
            entry.create_system = 3
            entry.external_attr = (0o600 & 0xFFFF) << 16
            target.writestr(entry, data, compress_type=zipfile.ZIP_DEFLATED, compresslevel=9)
    temporary.replace(path)


def font_path() -> str | None:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return None


def pil_font(size: int):
    fp = font_path()
    if fp:
        try:
            return ImageFont.truetype(fp, size)
        except Exception:
            pass
    return ImageFont.load_default()


def set_doc_defaults(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_header_footer(doc: Document, header: str):
    for section in doc.sections:
        section.header.paragraphs[0].text = header
        section.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        section.footer.paragraphs[0].text = "Confidential • Agentic Office UI browser verification"
        section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def make_docx_legal_contract(path: Path):
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "Master Services Agreement — Browser Verification Fixture")
    doc.add_heading("MASTER SERVICES AGREEMENT", 0)
    intro = doc.add_paragraph()
    intro.add_run("This Master Services Agreement ").bold = True
    intro.add_run("is entered into by and between Extend Labs Ltd. and Acme Manufacturing Group for a production-grade document rendering validation scenario.")

    clauses = [
        ("1. Scope of Services", "Provider will deliver document intelligence, extraction review, workflow automation, quality assurance dashboards, and support services for the customer’s finance and operations teams."),
        ("2. Service Levels", "Provider will maintain commercially reasonable availability targets, response windows, and escalation paths. Severity one incidents require acknowledgement within one business hour."),
        ("3. Data Processing", "Each party will process personal information in accordance with applicable law. Confidential data includes contracts, invoices, purchase orders, and employee records."),
        ("4. Fees and Payment", "Customer will pay monthly subscription fees, implementation charges, and usage-based overages according to the order form. Late payments may accrue interest."),
        ("5. Audit Rights", "Customer may audit security controls once per calendar year with reasonable notice. Provider may redact unrelated customer information from audit evidence."),
        ("6. Limitation of Liability", "Neither party will be liable for indirect, incidental, punitive, or consequential damages, subject to exclusions for confidentiality and data protection obligations."),
        ("7. Term and Termination", "The initial term is twenty-four months. Either party may terminate for uncured material breach after thirty days written notice."),
        ("8. Governing Law", "This agreement is governed by the laws of the State of New York, excluding conflict-of-law principles."),
    ]
    for i in range(3):
        for title, body in clauses:
            doc.add_heading(title, level=1)
            p = doc.add_paragraph(body)
            p.paragraph_format.space_after = Pt(6)
            for j in range(1, 4):
                doc.add_paragraph(f"{title}.{j} Operational note {j}: {body} This sentence is repeated to create realistic multi-page pagination and line wrapping behavior.", style="List Number")
        if i < 2:
            doc.add_page_break()

    doc.add_heading("Authorized Signatures", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Extend Labs Ltd."
    table.cell(0, 1).text = "Acme Manufacturing Group"
    table.cell(1, 0).text = "Name: Jordan Lee\nTitle: VP Operations\nDate: 2026-07-05"
    table.cell(1, 1).text = "Name: Morgan Chen\nTitle: CFO\nDate: 2026-07-05"
    table.cell(2, 0).text = "Signature: ____________________"
    table.cell(2, 1).text = "Signature: ____________________"
    doc.save(path)


def make_docx_invoice_table(path: Path):
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "Invoice Table Fixture")
    doc.add_heading("INVOICE INV-2026-0705", 0)
    doc.add_paragraph("Vendor: Northwind Industrial Supplies    Customer: Extend Labs Ltd.")
    doc.add_paragraph("Tax ID: US-98-7654321    Due Date: 2026-08-04")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Item", "Description", "Qty", "Unit Price", "Tax", "Line Total"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    rows = [
        ("A-100", "OCR document page processing", 1200, 0.18, 0.07),
        ("B-210", "Human-in-the-loop verification seats", 12, 89.00, 0.07),
        ("C-310", "Secure retention storage GB", 500, 0.12, 0.07),
        ("D-410", "Workflow automation connector", 3, 450.00, 0.07),
    ]
    subtotal = 0
    for item, desc, qty, price, tax in rows:
        line = qty * price
        subtotal += line
        cells = table.add_row().cells
        values = [item, desc, str(qty), f"${price:,.2f}", f"{tax:.0%}", f"${line:,.2f}"]
        for c, v in zip(cells, values):
            c.text = v
    merged = table.add_row().cells
    merged[0].merge(merged[4]).text = "Subtotal"
    merged[5].text = f"${subtotal:,.2f}"
    total = subtotal * 1.07
    merged2 = table.add_row().cells
    merged2[0].merge(merged2[4]).text = "Total Due"
    merged2[5].text = f"${total:,.2f}"
    doc.add_paragraph("Payment terms: Net 30. Please include invoice number on wire remittance.")
    doc.save(path)


def make_docx_report_with_image(path: Path, image_path: Path):
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "Quarterly Operations Report Fixture")
    doc.add_heading("Quarterly Operations Report", 0)
    p = doc.add_paragraph()
    p.add_run("Executive summary: ").bold = True
    p.add_run("Document automation throughput improved ").italic = True
    p.add_run("18.4% quarter over quarter").underline = True
    p.add_run(" while maintaining review accuracy above 99.2%.")
    r = doc.add_paragraph().add_run("Risk indicator: manual exception volume increased in the APAC queue.")
    r.font.color.rgb = RGBColor(180, 40, 40)
    doc.add_picture(str(image_path), width=Inches(5.2))
    doc.add_paragraph("Figure 1. Synthetic invoice image used to verify embedded image rendering.")
    doc.add_heading("Operational Highlights", level=1)
    for text in [
        "Average extraction latency fell from 4.2 seconds to 2.9 seconds.",
        "Reviewer acceptance rate stayed within the control band for three consecutive months.",
        "The top exception categories were missing purchase order, tax mismatch, and low-confidence vendor name.",
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(path)


def make_docx_chinese_mixed(path: Path):
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "中文混排测试文档")
    doc.add_heading("中文、English、数字与标点混排测试", 0)
    paragraphs = [
        "本文件用于验证 Vue DOCX Viewer 在中文、English words、1234567890、全角标点，半角 punctuation，以及长句换行场景下的渲染效果。",
        "合同编号：CN-EXT-2026-0705；供应商：上海凌云科技有限公司；金额：人民币 1,234,567.89 元；税率：6%。",
        "Mixed sentence: AI document processing 在真实业务中通常包含 invoices、contracts、receipts、forms 与 handwritten signatures。",
    ]
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.add_heading("编号列表", level=1)
    for i, text in enumerate(["准备测试物料", "启动浏览器验证", "记录样式与功能缺陷"], 1):
        doc.add_paragraph(f"{i}. {text}")
    doc.save(path)


def make_docx_review_comments(path: Path):
    """Create a deterministic review fixture with one comment and one replacement."""
    doc = Document()
    set_doc_defaults(doc)
    add_header_footer(doc, "Document Review Fixture")
    doc.add_heading("Supplier Review Notes", 0)
    doc.add_paragraph("COMMENT_TARGET")
    doc.add_paragraph("REVISION_TARGET")
    doc.add_paragraph("Use the 修订 and 批注 switches to compare the reviewed document state.")
    doc.save(path)

    with zipfile.ZipFile(path, "r") as source:
        entries = {name: source.read(name) for name in source.namelist()}

    document_xml = entries["word/document.xml"].decode("utf-8")
    comment_paragraph = (
        '<w:p><w:commentRangeStart w:id="0"/>'
        '<w:r><w:t>Monthly operational report</w:t></w:r>'
        '<w:commentRangeEnd w:id="0"/>'
        '<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
        '<w:commentReference w:id="0"/></w:r></w:p>'
    )
    revision_paragraph = (
        '<w:p><w:r><w:t xml:space="preserve">Response time changed from </w:t></w:r>'
        '<w:del w:id="1" w:author="Alex Reviewer" w:date="2026-07-11T00:00:00Z">'
        '<w:r><w:delText>five business days</w:delText></w:r></w:del>'
        '<w:r><w:t xml:space="preserve"> to </w:t></w:r>'
        '<w:ins w:id="2" w:author="Morgan Editor" w:date="2026-07-11T00:05:00Z">'
        '<w:r><w:t>two business days</w:t></w:r></w:ins><w:r><w:t>.</w:t></w:r></w:p>'
    )
    comment_paragraph = comment_paragraph.replace(
        "</w:p>",
        '<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:footnoteReference w:id="1"/></w:r></w:p>',
    )
    revision_paragraph = revision_paragraph.replace(
        "</w:p>",
        '<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:endnoteReference w:id="1"/></w:r></w:p>',
    )
    document_xml, comment_count = re.subn(
        r"<w:p\b[^>]*>(?:(?!</w:p>).)*?<w:t>COMMENT_TARGET</w:t>(?:(?!</w:p>).)*?</w:p>",
        comment_paragraph,
        document_xml,
        count=1,
    )
    document_xml, revision_count = re.subn(
        r"<w:p\b[^>]*>(?:(?!</w:p>).)*?<w:t>REVISION_TARGET</w:t>(?:(?!</w:p>).)*?</w:p>",
        revision_paragraph,
        document_xml,
        count=1,
    )
    if comment_count != 1 or revision_count != 1:
        raise RuntimeError("Unable to inject DOCX review fixture markup")
    entries["word/document.xml"] = document_xml.encode("utf-8")

    comments_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:comment w:id="0" w:author="Priya Reviewer" w:initials="PR" '
        'w:date="2026-07-11T00:10:00Z"><w:p><w:r>'
        '<w:t>Please confirm this report includes the service-level summary.</w:t>'
        '</w:r></w:p></w:comment></w:comments>'
    )
    entries["word/comments.xml"] = comments_xml.encode("utf-8")
    entries["word/footnotes.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:id="1"><w:p><w:r><w:t>脚注正文：服务级别摘要已经由审核人确认。</w:t></w:r></w:p></w:footnote>'
        '</w:footnotes>'
    ).encode("utf-8")
    entries["word/endnotes.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:endnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:endnote w:id="1"><w:p><w:r><w:t>尾注正文：响应时间变更将在下一版合同中生效。</w:t></w:r></w:p></w:endnote>'
        '</w:endnotes>'
    ).encode("utf-8")

    content_types = entries["[Content_Types].xml"].decode("utf-8")
    if "/word/comments.xml" not in content_types:
        content_types = content_types.replace(
            "</Types>",
            '<Override PartName="/word/comments.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
            "</Types>",
        )
    content_types = content_types.replace(
        "</Types>",
        '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
        '<Override PartName="/word/endnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"/>'
        "</Types>",
    )
    entries["[Content_Types].xml"] = content_types.encode("utf-8")

    relationships = entries["word/_rels/document.xml.rels"].decode("utf-8")
    if "relationships/comments" not in relationships:
        relationships = relationships.replace(
            "</Relationships>",
            '<Relationship Id="rIdReviewComments" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" '
            'Target="comments.xml"/></Relationships>',
        )
    relationships = relationships.replace(
        "</Relationships>",
        '<Relationship Id="rIdReviewFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
        '<Relationship Id="rIdReviewEndnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes" Target="endnotes.xml"/>'
        "</Relationships>",
    )
    entries["word/_rels/document.xml.rels"] = relationships.encode("utf-8")

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for name in sorted(entries):
            target.writestr(name, entries[name])


def make_invoice_png(path: Path):
    img = Image.new("RGB", (1000, 1330), "white")
    d = ImageDraw.Draw(img)
    title_font = pil_font(44)
    h_font = pil_font(26)
    font = pil_font(22)
    small = pil_font(18)
    d.rectangle([40, 40, 960, 1290], outline=(40, 60, 80), width=3)
    d.text((70, 70), "INVOICE", fill=(20, 50, 80), font=title_font)
    d.text((70, 135), "Northwind Industrial Supplies", fill=(0, 0, 0), font=h_font)
    d.text((70, 178), "Tax ID: US-98-7654321", fill=(0, 0, 0), font=font)
    d.text((650, 135), "Invoice #: INV-2026-0705", fill=(0, 0, 0), font=font)
    d.text((650, 178), "Date: 2026-07-05", fill=(0, 0, 0), font=font)
    d.text((70, 250), "Bill To: Extend Labs Ltd.", fill=(0, 0, 0), font=h_font)
    y = 340
    d.rectangle([70, y, 930, y+45], fill=(230, 237, 245), outline=(80, 100, 120))
    for x, t in [(90, "Description"), (560, "Qty"), (660, "Unit"), (800, "Total")]:
        d.text((x, y+10), t, fill=(0, 0, 0), font=small)
    y += 45
    items = [("OCR page processing", "1200", "$0.18", "$216.00"), ("Review seats", "12", "$89.00", "$1,068.00"), ("Storage GB", "500", "$0.12", "$60.00"), ("Workflow connector", "3", "$450.00", "$1,350.00")]
    for row in items:
        d.rectangle([70, y, 930, y+55], outline=(180, 190, 200))
        for x, t in [(90, row[0]), (570, row[1]), (660, row[2]), (800, row[3])]:
            d.text((x, y+15), t, fill=(0, 0, 0), font=small)
        y += 55
    d.text((650, y+45), "Subtotal", fill=(0, 0, 0), font=font)
    d.text((800, y+45), "$2,694.00", fill=(0, 0, 0), font=font)
    d.text((650, y+90), "Tax", fill=(0, 0, 0), font=font)
    d.text((800, y+90), "$188.58", fill=(0, 0, 0), font=font)
    d.text((650, y+145), "Total Due", fill=(180, 40, 40), font=h_font)
    d.text((800, y+145), "$2,882.58", fill=(180, 40, 40), font=h_font)
    d.text((70, 1180), "Payment terms: Net 30. Bank transfer preferred.", fill=(60, 60, 60), font=small)
    img.save(path)


def make_contract_page_png(path: Path):
    img = Image.new("RGB", (1000, 1330), "white")
    d = ImageDraw.Draw(img)
    title_font = pil_font(38)
    h_font = pil_font(26)
    font = pil_font(20)
    d.text((90, 70), "MASTER SERVICES AGREEMENT", fill=(20, 20, 20), font=title_font)
    y = 150
    sections = [
        ("1. Scope of Services", "Provider will deliver document intelligence, extraction review, workflow automation, and QA dashboards."),
        ("2. Service Levels", "Provider will maintain availability targets, response windows, and escalation paths for severity incidents."),
        ("3. Fees and Payment", "Customer will pay monthly subscription fees, implementation charges, and usage-based overages."),
    ]
    for heading, body in sections:
        d.text((90, y), heading, fill=(25, 65, 120), font=h_font)
        y += 46
        for line in wrap(body * 2, 82):
            d.text((90, y), line, fill=(35, 35, 35), font=font)
            y += 31
        y += 24
    d.rectangle([90, 870, 910, 1060], outline=(80, 80, 80), width=2)
    d.text((120, 900), "Authorized Signatures", fill=(20, 20, 20), font=h_font)
    d.line([120, 980, 450, 980], fill=(20, 20, 20), width=2)
    d.line([560, 980, 890, 980], fill=(20, 20, 20), width=2)
    d.text((120, 995), "Provider", fill=(35, 35, 35), font=font)
    d.text((560, 995), "Customer", fill=(35, 35, 35), font=font)
    img.save(path)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"], alignment=TA_CENTER, fontSize=24, leading=30, spaceAfter=16))
    return styles


def make_pdf_sample(path: Path):
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm)
    styles = pdf_styles()
    story = [Paragraph("Document Intelligence Verification Report", styles["CenterTitle"]), Paragraph("This text PDF contains searchable paragraphs, tables, and multiple pages for PdfViewer validation.", styles["BodyText"]), Spacer(1, 8)]
    data = [["Metric", "Jan", "Feb", "Mar"], ["Pages processed", "42,180", "47,920", "53,110"], ["Accuracy", "99.1%", "99.2%", "99.3%"], ["Exceptions", "813", "702", "655"]]
    t = Table(data, hAlign="LEFT")
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1C3A5E")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.4, colors.grey), ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F6FA")])]))
    story += [t, Spacer(1, 12)]
    for page in range(1, 5):
        story.append(Paragraph(f"Section {page}: Search Keyword ALPHA-{page}", styles["Heading1"]))
        for i in range(8):
            story.append(Paragraph("Agentic Office UI PdfViewer should support search, zoom, page navigation, thumbnail rendering, and scrolling with a realistic text layer. " * 2, styles["BodyText"]))
        if page < 4:
            story.append(PageBreak())
    doc.build(story)


def make_pdf_scanned(path: Path, image_path: Path):
    c = canvas.Canvas(str(path), pagesize=A4)
    w, h = A4
    image = ImageReader(str(image_path))
    for i in range(3):
        c.drawImage(image, 18*mm, 20*mm, width=w-36*mm, height=h-40*mm, preserveAspectRatio=True, anchor="c")
        c.drawString(20*mm, 10*mm, f"Scanned invoice page {i+1}")
        c.showPage()
    c.save()


def make_pdf_rotated(path: Path):
    c = canvas.Canvas(str(path), pagesize=A4)
    w, h = A4
    c.setFont("Helvetica-Bold", 22)
    c.drawString(40, h - 80, "Portrait Page")
    c.setFont("Helvetica", 12)
    for i in range(20):
        c.drawString(40, h - 120 - i*22, f"Portrait content row {i+1} for rotation and viewport validation.")
    c.showPage()
    c.setPageSize(landscape(A4))
    w, h = landscape(A4)
    c.setFont("Helvetica-Bold", 22)
    c.drawString(40, h - 70, "Landscape Page")
    for i in range(12):
        c.drawString(40, h - 110 - i*24, f"Landscape wide content row {i+1}: columns should remain visible during zoom and scroll testing.")
    c.showPage()
    c.save()


def make_pdf_large(path: Path):
    c = canvas.Canvas(str(path), pagesize=A4)
    w, h = A4
    for page in range(1, 32):
        c.setFont("Helvetica-Bold", 18)
        c.drawString(40, h - 60, f"Large Contract Page {page}")
        c.setFont("Helvetica", 10)
        for i in range(34):
            c.drawString(40, h - 95 - i*18, f"Clause {page}.{i+1}: This long contract page validates thumbnail buffering, scroll performance, and page counter updates.")
        c.showPage()
    c.save()


def style_header(ws):
    fill = PatternFill("solid", fgColor="1C3A5E")
    font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center")


def make_financial_model(path: Path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Assumptions"
    ws.append(["Driver", "FY2026", "FY2027", "FY2028"])
    drivers = [("Starting Customers", 120, 180, 255), ("ARPU", 420, 456, 492), ("Gross Margin", 0.72, 0.74, 0.76), ("Churn", 0.08, 0.07, 0.06)]
    for row in drivers:
        ws.append(row)
    style_header(ws)
    for col in range(1, 5):
        ws.column_dimensions[get_column_letter(col)].width = 18
    ws2 = wb.create_sheet("P&L")
    ws2.append(["Line Item", "FY2026", "FY2027", "FY2028"])
    rows = ["Revenue", "COGS", "Gross Profit", "Sales & Marketing", "R&D", "G&A", "Operating Income"]
    for r, label in enumerate(rows, 2):
        ws2.cell(r, 1, label)
    ws2["B2"] = "=Assumptions!B2*Assumptions!B3"
    ws2["C2"] = "=Assumptions!C2*Assumptions!C3"
    ws2["D2"] = "=Assumptions!D2*Assumptions!D3"
    for col in "BCD":
        ws2[f"{col}3"] = f"={col}2*(1-Assumptions!{col}4)"
        ws2[f"{col}4"] = f"={col}2-{col}3"
        ws2[f"{col}5"] = f"={col}2*22%"
        ws2[f"{col}6"] = f"={col}2*18%"
        ws2[f"{col}7"] = f"={col}2*11%"
        ws2[f"{col}8"] = f"={col}4-SUM({col}5:{col}7)"
    style_header(ws2)
    for row in ws2.iter_rows(min_row=2, min_col=2, max_col=4):
        for cell in row:
            cell.number_format = '$#,##0'
    chart = LineChart()
    chart.title = "Revenue Forecast"
    chart.y_axis.title = "USD"
    data = Reference(ws2, min_col=2, max_col=4, min_row=1, max_row=2)
    chart.add_data(data, titles_from_data=True)
    ws2.add_chart(chart, "F2")
    wb.create_sheet("Notes")["A1"] = "Financial model fixture with formulas and multiple sheets."
    wb.save(path)


def make_sales_table(path: Path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Sales Data"
    headers = ["Region", "Rep", "Product", "Month", "Units", "Unit Price", "Revenue", "Notes"]
    ws.append(headers)
    regions = ["North", "South", "East", "West", "APAC", "EMEA"]
    products = ["Doc Extract", "PDF Viewer", "XLSX Engine", "Workflow Bot"]
    for i in range(1, 181):
        units = random.randint(5, 180)
        price = random.choice([49, 89, 129, 249])
        ws.append([random.choice(regions), f"Rep {i%17+1}", random.choice(products), f"2026-{(i%12)+1:02d}", units, price, f"=E{i+1}*F{i+1}", "Long note validating overflow and cell text truncation behavior in the grid viewer."])
    style_header(ws)
    ws.freeze_panes = "A2"
    widths = [14, 12, 18, 12, 10, 12, 14, 58]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    for cell in ws[1]:
        cell.border = Border(bottom=Side(style="medium", color="1C3A5E"))
    sales_table = XlsxTable(displayName="SalesRecords", ref=f"A1:H{ws.max_row}")
    sales_table.tableStyleInfo = TableStyleInfo(
        name="TableStyleMedium2",
        showFirstColumn=False,
        showLastColumn=False,
        showRowStripes=True,
        showColumnStripes=False,
    )
    ws.add_table(sales_table)
    wb.save(path)


def make_charts_images(path: Path, image_path: Path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Dashboard"
    ws.append(["Quarter", "Revenue", "Exceptions"])
    for row in [("Q1", 120000, 820), ("Q2", 145000, 690), ("Q3", 168000, 540), ("Q4", 210000, 410)]:
        ws.append(row)
    style_header(ws)
    ws.merge_cells("A7:C8")
    ws["A7"] = "Worker merged region — image and merge parity"
    ws["A7"].font = Font(bold=True, color="1C3A5E")
    ws["A7"].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws["A2"].hyperlink = "#Dashboard!A7"
    ws["A2"].style = "Hyperlink"
    ws["B2"].comment = Comment("该季度收入已通过财务复核。", "财务审核")
    ws.conditional_formatting.add("B2:B5", ColorScaleRule(start_type="min", start_color="F8696B", mid_type="percentile", mid_value=50, mid_color="FFEB84", end_type="max", end_color="63BE7B"))
    ws.conditional_formatting.add("C2:C5", DataBarRule(start_type="min", end_type="max", color="5B9BD5", showValue=True))
    ws["D1"] = "趋势"
    ws["D1"].font = Font(color="FFFFFF", bold=True)
    ws["D3"], ws["D4"], ws["D5"] = 3, 2, 1
    ws.conditional_formatting.add("D3:D5", IconSetRule("3TrafficLights1", "num", [1, 2, 3], showValue=True))
    ws.row_dimensions[7].height = 28
    ws.row_dimensions[8].height = 28
    chart = BarChart()
    chart.title = "Revenue vs Quarter"
    chart.y_axis.title = "Revenue"
    data = Reference(ws, min_col=2, max_col=2, min_row=1, max_row=5)
    cats = Reference(ws, min_col=1, min_row=2, max_row=5)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    ws.add_chart(chart, "E2")
    chartsheet_chart = BarChart()
    chartsheet_chart.title = "Revenue by Quarter"
    chartsheet_chart.y_axis.title = "Revenue"
    chartsheet_chart.add_data(data, titles_from_data=True)
    chartsheet_chart.set_categories(cats)
    chart_sheet = wb.create_chartsheet("Revenue Chart")
    chart_sheet.add_chart(chartsheet_chart)
    img = XlsxImage(str(image_path))
    img.width = 240
    img.height = 320
    ws.add_image(img, "E18")
    wb.save(path)
    inject_xlsx_fidelity_parts(path)


def inject_xlsx_fidelity_parts(path: Path):
    with zipfile.ZipFile(path, "r") as source:
        entries = {name: source.read(name) for name in source.namelist()}

    sheet_path = "xl/worksheets/sheet1.xml"
    sheet_xml = entries[sheet_path].decode("utf-8")
    sparkline_xml = (
        '<extLst><ext uri="{05C60535-1F16-4fd2-B633-F4F36F0B64E0}">'
        '<x14:sparklineGroups xmlns:x14="http://schemas.microsoft.com/office/spreadsheetml/2009/9/main" '
        'xmlns:xm="http://schemas.microsoft.com/office/excel/2006/main">'
        '<x14:sparklineGroup type="line" markers="1"><x14:colorSeries rgb="FF2563EB"/>'
        '<x14:colorMarkers rgb="FFDC2626"/><x14:sparklines><x14:sparkline>'
        '<xm:f>Dashboard!B2:B5</xm:f><xm:sqref>D2</xm:sqref>'
        '</x14:sparkline></x14:sparklines></x14:sparklineGroup></x14:sparklineGroups>'
        '</ext></extLst>'
    )
    control_xml = (
        '<controls><control xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" shapeId="9001" name="发布前确认" r:id="rIdFidelityControl">'
        '<controlPr><xdr:anchor xmlns:xdr="http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing"><xdr:from><xdr:col>0</xdr:col><xdr:colOff>0</xdr:colOff><xdr:row>9</xdr:row><xdr:rowOff>0</xdr:rowOff></xdr:from>'
        '<xdr:to><xdr:col>2</xdr:col><xdr:colOff>0</xdr:colOff><xdr:row>11</xdr:row><xdr:rowOff>0</xdr:rowOff></xdr:to></xdr:anchor></controlPr>'
        '</control></controls>'
    )
    sheet_xml = sheet_xml.replace("</worksheet>", sparkline_xml + control_xml + "</worksheet>")
    entries[sheet_path] = sheet_xml.encode("utf-8")

    rels_path = "xl/worksheets/_rels/sheet1.xml.rels"
    rels_xml = entries[rels_path].decode("utf-8")
    rels_xml = rels_xml.replace(
        "</Relationships>",
        '<Relationship Id="rIdFidelityControl" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/ctrlProp" Target="../ctrlProps/ctrlProp1.xml"/>'
        "</Relationships>",
    )
    entries[rels_path] = rels_xml.encode("utf-8")
    entries["xl/ctrlProps/ctrlProp1.xml"] = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<formControlPr xmlns="http://schemas.microsoft.com/office/spreadsheetml/2009/9/main" objectType="CheckBox" checked="1" fmlaLink="Dashboard!$A$10"/>'
    ).encode("utf-8")

    drawing_path = "xl/drawings/drawing1.xml"
    drawing_xml = entries[drawing_path].decode("utf-8")
    shape_xml = (
        '<xdr:twoCellAnchor xmlns:xdr="http://schemas.openxmlformats.org/drawingml/2006/spreadsheetDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"><xdr:from><xdr:col>4</xdr:col><xdr:colOff>0</xdr:colOff><xdr:row>9</xdr:row><xdr:rowOff>0</xdr:rowOff></xdr:from>'
        '<xdr:to><xdr:col>7</xdr:col><xdr:colOff>0</xdr:colOff><xdr:row>12</xdr:row><xdr:rowOff>0</xdr:rowOff></xdr:to>'
        '<xdr:sp><xdr:nvSpPr><xdr:cNvPr id="9002" name="审批提醒" descr="只读形状展示"/><xdr:cNvSpPr/></xdr:nvSpPr>'
        '<xdr:spPr><a:solidFill><a:srgbClr val="DBEAFE"/></a:solidFill><a:ln w="19050"><a:solidFill><a:srgbClr val="2563EB"/></a:solidFill></a:ln><a:prstGeom prst="roundRect"><a:avLst/></a:prstGeom></xdr:spPr>'
        '<xdr:txBody><a:bodyPr anchor="ctr"/><a:lstStyle/><a:p><a:pPr algn="ctr"/><a:r><a:rPr lang="zh-CN" sz="1200" b="1"/><a:t>审批提醒：请核对季度收入</a:t></a:r></a:p></xdr:txBody>'
        '</xdr:sp><xdr:clientData/></xdr:twoCellAnchor>'
    )
    closing_tag = "</xdr:wsDr>" if "</xdr:wsDr>" in drawing_xml else "</wsDr>"
    drawing_xml = drawing_xml.replace(closing_tag, shape_xml + closing_tag)
    entries[drawing_path] = drawing_xml.encode("utf-8")

    content_types = entries["[Content_Types].xml"].decode("utf-8")
    content_types = content_types.replace(
        "</Types>",
        '<Override PartName="/xl/ctrlProps/ctrlProp1.xml" ContentType="application/vnd.ms-excel.controlproperties+xml"/>'
        "</Types>",
    )
    entries["[Content_Types].xml"] = content_types.encode("utf-8")

    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for name in sorted(entries):
            target.writestr(name, entries[name])


def make_large_grid(path: Path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Large Grid"
    headers = [f"Col {i}" for i in range(1, 61)]
    ws.append(headers)
    for r in range(1, 551):
        ws.append([f"R{r}C{c}" if c % 7 else r * c for c in range(1, 61)])
    style_header(ws)
    wb.save(path)


def make_json_fixtures():
    fields = {
        "file": "/samples/invoice.png",
        "fields": [
            {"id": "vendor", "label": "Vendor", "page": 1, "rect": [0.07, 0.10, 0.48, 0.15], "value": "Northwind Industrial Supplies", "confidence": 0.96},
            {"id": "tax_id", "label": "Tax ID", "page": 1, "rect": [0.07, 0.135, 0.33, 0.17], "value": "US-98-7654321", "confidence": 0.83},
            {"id": "invoice_no", "label": "Invoice #", "page": 1, "rect": [0.64, 0.10, 0.91, 0.14], "value": "INV-2026-0705", "confidence": 0.74},
            {"id": "total", "label": "Total Due", "page": 1, "rect": [0.64, 0.64, 0.92, 0.70], "value": "$2,882.58", "confidence": 0.42},
            {"id": "terms", "label": "Payment Terms", "page": 1, "rect": [0.07, 0.88, 0.58, 0.91], "value": "Net 30", "confidence": 0.91},
        ],
    }
    (SAMPLES / "field-citations.json").write_text(json.dumps(fields, ensure_ascii=False, indent=2), encoding="utf-8")

    layout = {
        "file": "/samples/contract-page.png",
        "output": {
            "width": 1000,
            "height": 1330,
            "blocks": [
                {"id": "title", "kind": "title", "bbox": [90, 65, 640, 55], "text": "MASTER SERVICES AGREEMENT", "confidence": 0.98},
                {"id": "h1", "kind": "header", "bbox": [90, 148, 360, 42], "text": "1. Scope of Services", "confidence": 0.93},
                {"id": "body1", "kind": "text", "bbox": [90, 195, 820, 130], "text": "Provider will deliver document intelligence and workflow automation.", "confidence": 0.89},
                {"id": "h2", "kind": "header", "bbox": [90, 360, 310, 42], "text": "2. Service Levels", "confidence": 0.91},
                {"id": "body2", "kind": "text", "bbox": [90, 410, 820, 130], "text": "Provider will maintain availability targets and escalation paths.", "confidence": 0.86},
                {"id": "table", "kind": "table", "bbox": [90, 870, 820, 190], "text": "Authorized Signatures", "confidence": 0.77},
                {"id": "footer", "kind": "footer", "bbox": [90, 1195, 650, 30], "text": "Confidential", "confidence": 0.64},
            ],
        },
    }
    (SAMPLES / "ocr-layout.json").write_text(json.dumps(layout, ensure_ascii=False, indent=2), encoding="utf-8")


def make_corrupted(path: Path, text: str):
    path.write_bytes((text * 20).encode("utf-8"))


def main():
    invoice_png = SAMPLES / "invoice.png"
    contract_png = SAMPLES / "contract-page.png"
    make_invoice_png(invoice_png)
    make_contract_page_png(contract_png)

    make_docx_legal_contract(SAMPLES / "legal-contract.docx")
    make_docx_invoice_table(SAMPLES / "invoice-table.docx")
    make_docx_report_with_image(SAMPLES / "report-with-image.docx", invoice_png)
    make_docx_chinese_mixed(SAMPLES / "chinese-mixed.docx")
    make_docx_review_comments(SAMPLES / "review-comments.docx")
    make_docx_legal_contract(SAMPLES / "demo.docx")
    make_corrupted(SAMPLES / "corrupted.docx", "This is not a valid DOCX package. ")

    make_pdf_sample(SAMPLES / "sample.pdf")
    make_pdf_scanned(SAMPLES / "scanned-invoice.pdf", invoice_png)
    make_pdf_rotated(SAMPLES / "rotated-pages.pdf")
    make_pdf_large(SAMPLES / "large-contract.pdf")
    make_corrupted(SAMPLES / "corrupted.pdf", "%PDF corrupted fixture\n")

    make_financial_model(SAMPLES / "financial-model.xlsx")
    make_sales_table(SAMPLES / "sales-table.xlsx")
    make_charts_images(SAMPLES / "charts-images.xlsx", invoice_png)
    make_large_grid(SAMPLES / "large-grid.xlsx")
    make_corrupted(SAMPLES / "corrupted.xlsx", "This is not a valid XLSX package. ")

    make_json_fixtures()

    for path in sorted([*SAMPLES.glob("*.docx"), *SAMPLES.glob("*.xlsx")]):
        if not path.name.startswith("corrupted."):
            normalize_ooxml_zip(path)

    manifest = []
    for path in sorted(SAMPLES.iterdir()):
        if path.is_file() and path.name != "manifest.json":
            manifest.append({"name": path.name, "bytes": path.stat().st_size})
    (SAMPLES / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
